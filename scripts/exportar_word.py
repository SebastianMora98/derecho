"""Convierte un documento de sección (markdown) a un .docx simple, para que se
pueda descargar y entregar en Word. No es un renderizador de markdown general:
solo entiende el subconjunto que usan estos documentos —párrafos, `## título`,
`**negrita**`, `*cursiva*` y listas `- ` (las referencias)—, que es a propósito
lo mismo que entiende `web/src/lib/md.ts` del lado del sitio.

El formato es el que suele pedirse en una entrega de Derecho: Times New Roman
12, doble espacio, márgenes de 2,54 cm, sangría de primera línea en el cuerpo,
sin colores ni estilos decorativos. Nada que no escribiría un estudiante desde
un Word en blanco.

Uso:
    uv run --with python-docx scripts/exportar_word.py entrada.md salida.docx --titulo "Título"
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


def curvar_comillas(texto: str) -> str:
    """"cita" → “cita”, como las deja el autocorregido de Word al escribir.

    Convierte solo pares completos. Una comilla suelta —abrir sin cerrar, señal
    de una cita cortada al copiar el fragmento— se deja recta a propósito: es
    mejor que quede visible el corte que curvarla mal y disimularlo.
    """

    def reemplazar(m: re.Match) -> str:
        return "“" + m.group(1) + "”"

    return re.sub(r'"([^"]*)"', reemplazar, texto)


def agregar_texto_inline(parrafo, texto: str) -> None:
    """Parte `**negrita**` y `*cursiva*` en runs. No anidan entre sí, igual que
    en los `.md` de origen."""
    texto = curvar_comillas(texto)
    patron = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for trozo in patron.split(texto):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            run = parrafo.add_run(trozo[2:-2])
            run.bold = True
        elif trozo.startswith("*") and trozo.endswith("*"):
            run = parrafo.add_run(trozo[1:-1])
            run.italic = True
        else:
            parrafo.add_run(trozo)


def sin_bordes_de_pagina(document: Document) -> None:
    """python-docx no expone `pgBorders`; sin este bloqueo, algunas plantillas
    de Word heredan un borde de página del `Normal.dotm` del usuario. No es el
    caso acá, pero lo forzamos a "ninguno" para no depender de eso."""
    sectPr = document.sections[0]._sectPr
    pgBorders = sectPr.makeelement(qn("w:pgBorders"), {})
    sectPr.append(pgBorders)


def construir(md_path: Path, salida: Path, titulo: str | None) -> None:
    texto = md_path.read_text(encoding="utf-8")
    lineas = texto.split("\n")

    doc = Document()

    seccion = doc.sections[0]
    seccion.page_width = Cm(21.59)  # carta, como el .docx del profesor
    seccion.page_height = Cm(27.94)
    for lado in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(seccion, lado, Cm(2.54))
    sin_bordes_de_pagina(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    if titulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        p.paragraph_format.space_after = Pt(12)

    # Agrupar líneas en bloques: cada bloque es un encabezado, un ítem de lista
    # o un párrafo. Una referencia ocupa varias líneas de fuente —la segunda en
    # adelante sangrada dos espacios—, así que solo una "- " SIN sangría abre un
    # ítem nuevo; cualquier otra línea, sangrada o no, sigue acumulando en el
    # bloque abierto hasta la próxima línea en blanco.
    bloques: list[tuple[str, str]] = []
    actual: list[str] = []
    tipo_actual = "parrafo"

    def cerrar():
        if actual:
            bloques.append((tipo_actual, " ".join(actual).strip()))
            actual.clear()

    for linea in lineas:
        cruda = linea.rstrip()
        if not cruda.strip():
            cerrar()
            tipo_actual = "parrafo"
            continue
        if cruda.startswith("## "):
            cerrar()
            tipo_actual = "parrafo"
            bloques.append(("titulo2", cruda[3:].strip()))
            continue
        item_nuevo = cruda.startswith("- ")
        if item_nuevo:
            cerrar()
            tipo_actual = "item"
        actual.append(cruda.lstrip()[2:] if item_nuevo else cruda.strip())
    cerrar()

    for tipo, contenido in bloques:
        if tipo == "titulo2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(contenido)
            run.bold = True
            continue
        if tipo == "item":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
            agregar_texto_inline(p, contenido)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        agregar_texto_inline(p, contenido)

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(salida))
    print(f"✓ {salida} ({len(bloques)} bloques)")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("entrada", type=Path)
    ap.add_argument("salida", type=Path)
    ap.add_argument("--titulo", default=None)
    args = ap.parse_args()
    construir(args.entrada, args.salida, args.titulo)
